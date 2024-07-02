"""
Utilities used for Data Evaluation

This class contains the following classes:

- DataAnalysis: performs data analysis for initial evaluation of data
- DataQualityReport: Generates a data quality report based on analysis performed by DataAnalysis. 
"""

import pandas as pd
import dataframe_image as dfi
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


class DataAnalysis:
    """Class for performing analysis on a given dataframe
    """

    def __init__(self, name: str, df: pd.DataFrame):
        """Initialize DataAnalysis

        Args:
            name (str): Name of the data to be analysed - used for naming of images and csv files.
            df (pd.DataFrame): DataFrame to be analysed.
        """

        self.name = name
        self.df = df
        # File names for descriptive feature tables - allows tables to be inserted into data quality report
        self.numeric_table_filename = f"{self.name}-DataQualityReport-NumericFeatures-Table.png"
        self.category_table_filename = f"{self.name}-DataQualityReport-CategoricFeatures-Table.png"
        # Initialize as empty list, add filenames when histograms added
        self.histogram_filenames = []
        self.boxplot_filenames = []  # Initialize as empty list.

        # Perform initial analysis on df
        self.initial_analysis()

    def get_numeric_columns(self):
        """Get numeric columns from the dataframe.
        """
        self.numeric_columns = self.df.select_dtypes(
            ['int64', 'float64', 'datetime64']).columns
        self.num_numeric_columns = len(self.numeric_columns)

    def get_category_columns(self):
        """Get categoric columns from the dataframe.
        """
        self.category_columns = self.df.select_dtypes(['category']).columns
        self.num_category_columns = len(self.category_columns)

    def get_cardinalities(self):
        """Caluculate the cardinalities of all numeric features on the DataFrame
        """
        # Check if numeric columns defined
        if self.numeric_columns is None:
            self.get_numeric_columns()

        # Get cardinality of the numeric columns
        self.numeric_columns_card = self.df[self.numeric_columns].nunique()
        self.df_numeric_card = pd.DataFrame(
            self.numeric_columns_card, columns=['card'])

    def get_shape(self):
        """Print the shape of the dataframe.
        """
        # Outputs # rows
        print("Number of rows:", self.df.shape[0])
        # Outputs # cols
        print("Number of cols:", self.df.shape[1])
        self.num_rows = self.df.shape[0]
        self.num_cols = self.df.shape[1]

    def print_col_types(self):
        """Get the column types of the DataFrame.
        """
        print("\n\n")
        print("Column types:")
        for col in self.df.columns:
            print(f"{col}: {self.df[col].dtype}")

    def convert_object_cols(self):
        """Convert all of the object cols to category
        """
        object_columns = self.df.select_dtypes(['object']).columns
        for column in object_columns:
            self.df[column] = self.df[column].astype('category')

    def count_duplicates(self):
        """Calculate number of duplicate rows"""
        self.duplicate_count = self.df[self.df.duplicated() == True].shape[0]
        print("Number of duplicate rows:", self.duplicate_count)

    def get_null_values(self):
        """Count null values per column in dataframe.
        """
        self.num_null_vals = {}

        for col in self.df.columns:
            self.num_null_vals[col] = self.df[col].isnull().sum()
            print(col, self.df[col].isnull().sum())

    def initial_analysis(self):
        """Perform initial analysis on the DataFrame.
        Allows for type checking before bulk analysis is performed.
        """
        self.get_shape()
        self.df.head()

        self.print_col_types()
        print("\n\n Convert column types to desired ones before continuing")

    def save_df_to_png(self, df, name):
        """Saves DataFrame to PNG image

        Args:
            df (pd.DataFrame): Dataframe to save
            name (str): File name to save to
        """
        dfi.export(df, name)

    def numeric_desc_features(self):
        """Create a numeric descriptive feature table and save it to a csv and png file.
        """
        # First describe the numeric features
        self.df_table_numeric = self.df[self.numeric_columns].describe().T

        # Add % missing
        numeric_columns_missing = 100 * \
            (self.df[self.numeric_columns].isnull().sum()/self.num_rows)
        self.df_numeric_missing = pd.DataFrame(
            numeric_columns_missing, columns=['%missing'])

        # Concatenate numeric, missing and cardinality cols
        df_numeric_columns_data_quality_report_table = pd.concat(
            [self.df_table_numeric, self.df_numeric_missing, self.df_numeric_card], axis=1)

        # Print data quality report table for numeric features to a file.
        df_numeric_columns_data_quality_report_table.to_csv(f"{self.name}-Amenities-DataQualityReport-NumericFeatures-Table.csv",
                                                            index_label='Feature')
        self.save_df_to_png(df_numeric_columns_data_quality_report_table,
                            self.numeric_table_filename)
        print(df_numeric_columns_data_quality_report_table)

    def categoric_desc_features(self):
        """Create a categoric descriptive feature table and save it to a csv and png file.
        """
        self.category_unique = self.df[self.category_columns].nunique()
        # Look once again at the summary stats table for categorical features
        self.df_table_categoric = self.df[self.category_columns].describe(
        ).T
        category_columns_perc_missing = 100 * \
            (self.df[self.category_columns].isnull().sum()/self.df.shape[0])
        self.df_category_perc_missing = pd.DataFrame(
            category_columns_perc_missing, columns=['%missing'])
        # cardinality
        category_columns_card = self.df[self.category_columns].nunique()
        self.df_category_card = pd.DataFrame(
            category_columns_card, columns=['card'])

        second_val = []
        second_freq = []
        for category in self.category_columns:
            # print(df[category].value_counts().index.tolist())
            if len(self.df[category].value_counts().index.tolist()) > 1:
                second_val.append(
                    self.df[category].value_counts().index.tolist()[1])
                second_freq.append(
                    self.df[category].value_counts().iloc[1])
            else:
                second_val.append("None")
                second_freq.append(0)

        # Concatenate numeric, missing and cardinality cols
        df_categoric_columns_data_quality_report_table = pd.concat(
            [self.df_table_categoric, self.df_category_perc_missing, self.df_category_card], axis=1)

        df_categoric_columns_data_quality_report_table['second'] = second_val
        df_categoric_columns_data_quality_report_table['second_freq'] = second_freq

        df_categoric_columns_data_quality_report_table = df_categoric_columns_data_quality_report_table[[
            'count', 'unique', 'top', 'freq', 'second', 'second_freq', '%missing', 'card']]
        # Print data quality report table for numeric features to a file.
        df_categoric_columns_data_quality_report_table.to_csv(f"{self.name}-DataQualityReport-CategoricFeatures-Table.csv",
                                                              index_label='Feature')
        self.save_df_to_png(df_categoric_columns_data_quality_report_table,
                            self.category_table_filename)

    def print_value_proportions(self):
        """Print the unique values and their corresponding proportions for each categoric column in the dataframe.
        """
        print("\n\nValue proportions:")
        # Change category cols to remove
        # Look at the values taken by each categorical feature, as a proportion, including NaN
        for column in self.category_columns:
            # print("\n" + column)
            print(self.df[column].value_counts(normalize=True, dropna=False))

    def plot_numeric(self):
        """Plot histogram for all numeric features.
        """
        # TODO: Make new fig every 10 plots for dataframes with large num of desc features
        # Plot histogram
        plt.figure()
        self.df.hist(figsize=(30, 20))

        # Wipe histogram filenames
        self.histogram_filenames = []

        # Save histogram to png and append to filenames
        f = f"{self.name}_numeric_hist.png"
        self.histogram_filenames.append(f)
        plt.savefig(f)

    def plot_categoric(self):
        """Create box plots for all categoric columns and save to png
        """
        num_cols = 2  # number of plots per row
        # Only want 10 plots per figure as image becomes too large for data quality report otherwise
        plots_per_fig = 10

        num_figures = (len(self.category_columns) +
                       plots_per_fig - 1)//plots_per_fig

        self.boxplot_filenames = []

        # Only add 10 plots to each figure, need to iterate num_figures times
        for i in range(num_figures):
            start = i * plots_per_fig
            end = min(
                len(self.category_columns), start + plots_per_fig)

            # Columns to plot in this iteration of the for loop
            current_cols = self.category_columns[start: end]
            # Need to calculate in for loop as last iteration will be different
            num_rows = (len(current_cols) + 1) // num_cols

            # Create a single figure with subplots arranged in a 2-column grid
            fig, axes = plt.subplots(
                nrows=num_rows, ncols=num_cols, figsize=(15, 5*num_rows))

            # Flatten the axes array to simplify indexing
            axes = axes.flatten()

            # Iterate over each column
            for j, column in enumerate(current_cols):
                # Plot each bar plot on its corresponding subplot
                self.df[column].value_counts(
                    dropna=False).plot(kind='bar', ax=axes[j])
                axes[j].set_title(column)

            for j in range(len(current_cols), num_rows * 2):
                fig.delaxes(axes[j])

            f = f"{self.name}_categoric_box_plot_{i}.png"
            self.boxplot_filenames.append(f)
            plt.savefig(f)
            plt.show()

    def analyse(self, plots=True):
        """Performs analysis on the dataframe after initial cleaning and column type conversion is done.

        Args:
            plots (bool, optional): Whether or not to draw plots of the features. Defaults to True.
        """
        self.get_numeric_columns()

        self.convert_object_cols()

        self.get_category_columns()

        self.count_duplicates()

        self.get_null_values()

        self.get_numeric_columns()

        self.get_cardinalities()

        self.print_value_proportions()

        self.numeric_desc_features()

        self.categoric_desc_features()

        if plots:
            self.plot_numeric()

            self.plot_categoric()

    def make_data_quality_report(self, params):
        self.dqr = DataQualityReport(self, params)


class DataQualityReport:
    critical_params = ['title', 'description',
                       'source', 'source_link', 'detailed_desc']
    sig_digits = 5
    drop_feature_cutoff = 50  # % of missing values allowed before feature is dropped

    def __init__(self, da: DataAnalysis, params: dict):
        self.da = da
        self.params = params
        # make sure all critical parameters are present
        for param in self.critical_params:
            if param not in params:
                raise ValueError(
                    f"The required value {param} is missing from parameters")

        # Get the columns that contain null values
        self.null_val_cols = []
        for col in self.da.num_null_vals:
            if self.da.num_null_vals[col] > 0:
                self.null_val_cols.append(col)

    def num_format(self, num):
        return f'%.{self.sig_digits}g' % num

    def write_missing_vals_overview(self):

        if len(self.null_val_cols) == 0:
            return "The dataset contains no missing values. "

        if len(self.null_val_cols) == 1:
            return f"The feature {self.null_val_cols[0]} contains missing values. "

        if len(self.null_val_cols) > 5:
            return f"{len(self.null_val_cols)} features contain missing values. "

        output = "The features "
        for col in self.null_val_cols:
            output += f"{col}, "

        output = output[:-2] + " contain missing values."

    def write_df_shape(self):
        return f"The dataset has {self.da.num_cols} features and {self.da.num_rows} rows. "

    def write_missing_values_summary(self):
        dropped_features = []
        for col in self.null_val_cols:
            if col in self.da.category_columns:
                missing = self.da.df_category_perc_missing.loc[col, '%missing']
            else:
                missing = self.da.df_numeric_missing.loc[col, '%missing']

            if col in self.da.category_columns and missing > 50:
                dropped_features.append(col)

        if len(dropped_features) == 0:
            return "No features will have to be dropped due to missing values. "

        if len(dropped_features) < 5:
            output = ""
            for feature in dropped_features:
                output += f"{feature}, "
            output = output[:-2] + \
                " will have to be dropped due to missing values. "
            return output

        return f"{len(dropped_features)} features will have to be dropped due to missing values. "

    def write_duplicate_rows(self):
        if (self.da.duplicate_count == 0):
            return "There are no duplicate rows. "

        return f"There are {self.da.duplicate_count} duplicate rows. "

    def numeric_col_description(self, col):
        output = f"This feature has a mean of {self.num_format(self.da.df_table_numeric.loc[col, 'mean'])}, a min value of {self.num_format(self.da.df_table_numeric.loc[col, 'min'])} and a max value of {self.num_format(self.da.df_table_numeric.loc[col, 'max'])}. "
        if col in self.null_val_cols:
            output += f"There are {self.da.num_null_vals[col]} missing values. "
        else:
            output += "There are no missing values. "

        return output

    def category_col_description(self, col):
        output = f"This has {self.da.category_unique[col]} unique values. The most common is {self.da.df_table_categoric.loc[col, 'top']}. "
        if col in self.null_val_cols:
            output += f"There are {self.da.num_null_vals[col]} missing values. "
        else:
            output += "There are no missing values. "

        return output

    def write_actions(self, document):

        # Get rows to drop for missing vals
        drop_row_cols = []
        for col in self.critical_params:
            if col in self.null_val_cols:
                drop_row_cols.append(col)

        cols_to_drop = self.da.df_category_perc_missing.loc[
            self.da.df_category_perc_missing['%missing'] > 50].index.values

        document.add_heading("Actions to Take", 2)
        document.add_paragraph(f"{len(cols_to_drop) + len(drop_row_cols)} actions will be taken:")
        for col in drop_row_cols:
            document.add_paragraph(col, style="List Bullet")
            document.add_paragraph(f"Drop rows with missing {col}. ", style="List Bullet 2")

        for col in cols_to_drop:
            document.add_paragraph(col, style="List Bullet")
            document.add_paragraph(f"Drop feature due to {self.num_format(self.da.df_category_perc_missing.loc[col, '%missing'])}% of values missing. ", style="List Bullet 2")

        p = document.add_paragraph("")
        p.add_run("***ADD ACTIONS TO TAKE***").bold = True

    def add_appendix(self, document):
        document.add_heading("Appendix", 1)
        document.add_heading("Continuous Features", 2)
        document.add_heading("Descriptive Statistics", 3)
        document.add_picture(self.da.numeric_table_filename, width=Inches(6.5))
        document.add_heading("Histograms", 3)
        for f in self.da.histogram_filenames:
            document.add_picture(f, width=Inches(6.5))

        document.add_heading("Categorical Features", 2)
        document.add_heading("Descriptive Statistics", 3)
        document.add_picture(self.da.category_table_filename, width=Inches(6.5))
        document.add_heading("Box Plots", 3)
        for f in self.da.boxplot_filenames:
            document.add_picture(f, width=Inches(5))

    def write_document(self):
        document = Document()

        document.add_heading(f"{self.params['title']} Data Quality Report", 1)
        document.add_heading('Overview', 2)

        p = document.add_paragraph(f"This report will outline the initial data quality findings on {self.params['description']} data obtained from {self.params['source']} which can be found at {self.params['source_link']}. ")
        p.add_run(f'This report will include an overview of the dataset, and a review of the continuous and categorical features, including histograms and bar charts. On initial review, this dataset contains a lot of missing data for most features. The data that is present appears to be reasonable and logical, however a number of columns will need to be dropped. ')
        mv = self.write_missing_vals_overview()
        p.add_run(mv)

        document.add_heading('Summary', 2)
        p = document.add_paragraph(f"This dataset consists of {self.params['detailed_desc']}. ")
        p.add_run(self.write_df_shape())
        p.add_run(self.write_missing_values_summary())
        p.add_run(self.write_duplicate_rows())
        p.add_run("Distribution of the data is consistent with expectations.")

        document.add_heading('Review Logical Integrity', 2)
        document.add_paragraph('Test 1: ')
        document.add_paragraph('x instances.', style='List Bullet')

        if len(self.da.numeric_columns) > 0:
            document.add_heading("Review Continuous Features", 2)
            document.add_paragraph(f"There are {len(self.da.numeric_columns)} continuous features in this dataset:")
            for col in self.da.numeric_columns:
                document.add_paragraph(col, style="List Bullet")
                desc = self.numeric_col_description(col)
                document.add_paragraph(desc, style="List Bullet 2")

            document.add_heading("Histograms", 3)
            p = document.add_paragraph(
                "All Histograms can be found in the appendix as a summary sheet. All features show a plausible distribution.")
            p.add_run("***REVIEW DISTRIBUTION***").bold = True

        if len(self.da.category_columns) > 0:
            document.add_heading("Review Categoric Features", 2)
            document.add_paragraph(f"There are {len(self.da.category_columns)} categoric features in this dataset:")
            for col in self.da.category_columns:
                document.add_paragraph(col, style="List Bullet")
                desc = self.category_col_description(col)
                document.add_paragraph(desc, style="List Bullet 2")

        self.write_actions(document)

        self.add_appendix(document)

        document.save(f"{self.params['title']} Data Quality Report.docx")
